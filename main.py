#Authored by Seaghan Ennis
#Last Update: Jan 13, 2022

import docx
import re

def main():
    #enter script
    print("Starting formatter")

    #opens the document and snapshots it
    inputFile = "example.docx"
    print("Opening: ", inputFile)
    inDoc = docx.Document(inputFile)

    #setup outDoc
    outDoc = docx.Document()
    style = outDoc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = docx.shared.Pt(12)
    paragraph_format = outDoc.styles['Normal'].paragraph_format
    paragraph_format.space_before = docx.shared.Pt(0)
    paragraph_format.space_after = docx.shared.Pt(0)
    paragraph_format.line_spacing = 1

    #to calculate the total
    total = 0.0

    #Read inDoc to produce outDoc
    for p in inDoc.paragraphs:

        # Remove space
        for char in p.text:
            if char == " ":
                p.text = p.text[1:]
            else:
                break
        #find subtotals and add them to total
        if "Subtotal" in p.text:
            findSub = re.sub("[^0-9]", "", p.text)
            findSub = float(findSub)
            findSub = findSub * 0.01
            print("Found subtotal! Adding", findSub, "to full total")
            total = total + findSub

        #add all formated paragraphs to outDoc
        outDoc.add_paragraph(p.text)

    #Calculate totals
    discount = round(total * 0.15, 2)
    discountedTotal = total - discount

    #Add final information to doc
    outDoc.add_page_break()
    para = outDoc.add_paragraph("Total:\t\t\t\t\t")
    para.add_run('$' + str(total))
    para = outDoc.add_paragraph("15% Discount:\t\t\t\t")
    para.add_run('$(' + str(discount) + ')').underline = True
    para = outDoc.add_paragraph('')
    para.add_run("TOTAL:\t\t\t\t$" + str(discountedTotal)).bold = True

    outDoc.save("exampleOutput.docx")

    print("Finished and saved")

if __name__ == '__main__':
    #start script
    main()